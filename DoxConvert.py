import docx

#Crear los objetos que representan a cada archivo
doc1 = docx.Document("Telefonica Central America VPN Service FRS v0.1.docx")
doc2 = docx.Document("FRS_Telefonica Central America VPN Service FRS v0.1.docx")

#Leer el contenido completo de cada archivo
parrafo_doc1 = doc1.paragraphs
parrafo_doc2 = doc2.paragraphs

#Obtner el contenido de cada archivo en formato de texto
separador =  '\n' + 50*'-' + '\n'
archivo1 = [parrafo.text + separador for parrafo in parrafo_doc1]
archivo1 = ''.join(archivo1)

archivo2 = [parrafo.text + separador for parrafo in parrafo_doc2]
archivo2 = ''.join(archivo2)

with open('Telefonica Central America VPN Service FRS v0.1.txt', 'w') as file1:
    file1.write(archivo1)

with open('FRS_Telefonica Central America VPN Service FRS v0.1.txt', 'w') as file2:
    file2.write(archivo2)
